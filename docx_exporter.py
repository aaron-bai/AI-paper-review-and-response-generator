"""
docx_exporter.py
Export review questions and author responses to academic-style Word documents.
"""

from __future__ import annotations

import os
import re
from datetime import datetime
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


_ITEM_PATTERN = re.compile(r"^\s*(\d+)\.\s", re.MULTILINE)
_QUESTION_PREFIX_PATTERN = re.compile(r"^\s*\[Question\s*\d+\]\s*", re.IGNORECASE)


def _set_run_font(run, latin_font: str, east_asia_font: str, size: int, bold: bool = False):
    run.font.name = latin_font
    run.font.size = Pt(size)
    run.bold = bold
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:eastAsia"), east_asia_font)


def _set_paragraph_format(paragraph, line_spacing: float = 1.5, first_line_chars: float | None = 2.0):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_after = Pt(6)
    if first_line_chars is not None:
        fmt.first_line_indent = Cm(0.74 * first_line_chars)


def _set_page_layout(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def _configure_document_defaults(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    run_properties = normal._element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:eastAsia"), "宋体")
    _set_page_layout(document)


def _add_reference_runs(paragraph, reference_text: str):
    marker = paragraph.add_run("[引用定位] ")
    _set_run_font(marker, "Times New Roman", "黑体", 12, bold=True)
    marker.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    reference = paragraph.add_run(reference_text.strip())
    _set_run_font(reference, "Times New Roman", "宋体", 12)
    reference.font.italic = True


def _split_reference(item_text: str) -> tuple[str | None, str]:
    if not item_text.startswith("[Reference:"):
        return None, item_text.strip()

    closing = item_text.find("]")
    if closing == -1:
        return None, item_text.strip()

    reference_text = item_text[len("[Reference:"):closing].strip()
    body = item_text[closing + 1:].strip()
    return reference_text, body


def _split_numbered_items(text: str) -> list[str]:
    matches = list(_ITEM_PATTERN.finditer(text))
    if not matches:
        cleaned = text.strip()
        return [cleaned] if cleaned else []

    items: list[str] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        items.append(text[start:end].strip())
    return items


def _add_metadata(document: Document, paper_path: str, discipline: str, generated_at: datetime):
    metadata_lines = (
        ("论文文件", paper_path),
        ("学科领域", discipline),
        ("生成时间", generated_at.strftime("%Y-%m-%d %H:%M:%S")),
    )
    for label, value in metadata_lines:
        paragraph = document.add_paragraph()
        _set_paragraph_format(paragraph, first_line_chars=None)
        run = paragraph.add_run(f"{label}: {value}")
        _set_run_font(run, "Times New Roman", "宋体", 12)


def _add_title(document: Document, title: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    _set_run_font(run, "Times New Roman", "黑体", 16, bold=True)


def _add_section_heading(document: Document, heading: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(heading)
    _set_run_font(run, "Times New Roman", "黑体", 14, bold=True)


def _add_body_runs(paragraph, text: str):
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run(line)
        _set_run_font(run, "Times New Roman", "宋体", 12)


def _write_review_items(document: Document, items: Iterable[str]):
    for index, item in enumerate(items, start=1):
        reference_text, body = _split_reference(item)

        number_paragraph = document.add_paragraph()
        _set_paragraph_format(number_paragraph, first_line_chars=None)
        number_paragraph.paragraph_format.left_indent = Cm(0)
        number_run = number_paragraph.add_run(f"{index}.")
        _set_run_font(number_run, "Times New Roman", "黑体", 12, bold=True)

        if reference_text:
            reference_paragraph = document.add_paragraph()
            _set_paragraph_format(reference_paragraph, first_line_chars=None)
            reference_paragraph.paragraph_format.left_indent = Cm(0.74)
            _add_reference_runs(reference_paragraph, reference_text)

        body_paragraph = document.add_paragraph()
        _set_paragraph_format(body_paragraph)
        body_paragraph.paragraph_format.left_indent = Cm(0.74)
        _add_body_runs(body_paragraph, body)


def _write_response_items(document: Document, items: Iterable[str]):
    for index, item in enumerate(items, start=1):
        cleaned_item = _QUESTION_PREFIX_PATTERN.sub("", item).strip()

        question_paragraph = document.add_paragraph()
        _set_paragraph_format(question_paragraph, first_line_chars=None)
        title_run = question_paragraph.add_run(f"问题 {index}")
        _set_run_font(title_run, "Times New Roman", "黑体", 12, bold=True)

        explanation_labels = ("Explanation:", "问题说明：", "说明：")
        resolution_labels = ("Resolution:", "拟采取的修改与回复：", "回复：", "解决方案：")
        explanation_start, explanation_label = _find_first_label(cleaned_item, explanation_labels)
        resolution_start, resolution_label = _find_first_label(cleaned_item, resolution_labels)

        explanation_text = cleaned_item
        resolution_text = ""
        if explanation_start != -1 and resolution_start != -1:
            explanation_text = cleaned_item[
                explanation_start + len(explanation_label):resolution_start
            ].strip()
            resolution_text = cleaned_item[resolution_start + len(resolution_label):].strip()

        if explanation_text:
            explanation_paragraph = document.add_paragraph()
            _set_paragraph_format(explanation_paragraph)
            explanation_paragraph.paragraph_format.left_indent = Cm(0.74)
            label_run = explanation_paragraph.add_run("问题说明：")
            _set_run_font(label_run, "Times New Roman", "黑体", 12, bold=True)
            _add_body_runs(explanation_paragraph, explanation_text)

        if resolution_text:
            resolution_paragraph = document.add_paragraph()
            _set_paragraph_format(resolution_paragraph)
            resolution_paragraph.paragraph_format.left_indent = Cm(0.74)
            label_run = resolution_paragraph.add_run("拟采取的修改与回复：")
            _set_run_font(label_run, "Times New Roman", "黑体", 12, bold=True)
            _add_body_runs(resolution_paragraph, resolution_text)


def _add_footer_note(document: Document, note: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(note)
    _set_run_font(run, "Times New Roman", "宋体", 10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _find_first_label(text: str, labels: tuple[str, ...]) -> tuple[int, str]:
    best_index = -1
    best_label = ""
    for label in labels:
        index = text.find(label)
        if index != -1 and (best_index == -1 or index < best_index):
            best_index = index
            best_label = label
    return best_index, best_label


def export_review_documents(
    output_dir: str,
    paper_path: str,
    discipline: str,
    review_questions: str,
    review_responses: str,
    generated_at: datetime | None = None,
) -> tuple[str, str]:
    """Export review questions and responses to two docx files."""
    try:
        os.makedirs(output_dir, exist_ok=True)
    except OSError as exc:
        raise RuntimeError(f"无法创建输出目录: {output_dir}") from exc

    timestamp = (generated_at or datetime.now()).strftime("%Y%m%d_%H%M%S")
    review_path = os.path.join(output_dir, f"评审意见+{timestamp}.docx")
    response_path = os.path.join(output_dir, f"回复+{timestamp}.docx")

    review_doc = Document()
    _configure_document_defaults(review_doc)
    _add_title(review_doc, "学术论文评审意见")
    _add_metadata(review_doc, paper_path, discipline, generated_at or datetime.now())
    _add_section_heading(review_doc, "评审意见正文")
    _write_review_items(review_doc, _split_numbered_items(review_questions))
    _add_footer_note(review_doc, "注：引用定位内容直接保留原文标注，便于与论文正文逐项核对。")
    review_doc.save(review_path)

    response_doc = Document()
    _configure_document_defaults(response_doc)
    _add_title(response_doc, "学术论文评审回复")
    _add_metadata(response_doc, paper_path, discipline, generated_at or datetime.now())
    _add_section_heading(response_doc, "作者回复正文")
    _write_response_items(response_doc, _split_numbered_items(review_responses))
    _add_footer_note(response_doc, "注：回复按问题编号组织，建议结合修订稿逐条对应落实。")
    response_doc.save(response_path)

    return review_path, response_path